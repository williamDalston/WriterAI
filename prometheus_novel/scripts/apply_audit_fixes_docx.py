"""Apply audit fixes to a .docx manuscript (Burning Vows Final).

Fixes: ellipsis breakage, block duplication, Marco's wedding ring, grounding tics.
Deterministic, no LLM. Run on Burning Vows Final_KDP.docx or similar.

Usage:
    python -m prometheus_novel.scripts.apply_audit_fixes_docx [path/to/manuscript.docx]
"""

import re
import sys
from pathlib import Path

if sys.platform == "win32":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
except ImportError:
    print("[ERROR] pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = PROJECT_ROOT / "data/projects/burning-vows-30k/output/Burning Vows Final_KDP.docx"

# Grammarly ellipsis breakage: . -> ... in hesitation dialogue
ELLIPSIS_FIXES = [
    (r"Mamma gets\.\s+She", r"Mamma gets... She"),
    (r"has been\.\s+evolving", r"has been... evolving"),
    (r"relatives have\.\s+opinions", r"relatives have... opinions"),
    (r"Also\.\s+maybe", r"Also... maybe"),
    (r"Just\.\s+don", r"Just... don"),
    (r"Voglio\.\s+essere", r"Voglio... essere"),
    (r"about the\.\s+distraction", r"about the... distraction"),
    (r"you two\.\s+talking", r"you two... talking"),
    (r"name and thought\.\s+later", r"name and thought... later"),
]

# Marco's wedding ring (he's not married - remove references)
MARCO_RING_FIXES = [
    (r"The wedding ring on her finger is different from Marco's[—\-–]\s*older", r"The wedding ring on her finger is older"),
    (r"different from Marco's[—\-–]\s*older", r"older"),
]

# "I pressed my temple" grounding tic - remove clause, keep 2 (remove from others)
TEMPLE_REMOVALS = [
    (r"\s*I pressed my temple\.\s*", " "),
    (r" My jaw clenched as I pressed my temple\.", r" My jaw clenched."),
    (r"I pressed my temple and slowly turned", r"I slowly turned"),
    (r"I pressed my temple and gripped", r"I gripped"),
    (r"\s*I pressed my temple, feeling[^.]*?\.\s*", " "),
    (r"\s*I pressed my temple against[^.]*?\.\s*", " "),
]

# "leaving it ruffled/rumpled" - diversify
RUFFLED_FIXES = [
    (r"leaving it rumpled in a way that made him look younger", r"his hair mussed from the wind"),
    (r"leaving it ruffled\.", r"his hand still on the rail."),
]

# "like a confession" / "like a verdict" - reduce overuse
CONFESSION_VERDICT_FIXES = [
    (r"holding heat like a confession", r"holding heat"),
    (r"like a verdict\.", r"."),
]

# POV pronoun: "my" -> "her" when describing Sofia/Gianna (audit L20, L1144, L1228)
POV_FIXES = [
    (r", my voice taking", r", her voice taking"),
    (r"My voice carries absolute certainty", r"Her voice carries absolute certainty"),
    (r"my fingers brushing mine", r"her fingers brushing mine"),
    (r"My eyes find mine through dark lashes", r"Her eyes find mine through dark lashes"),
]


def apply_text_fixes(text: str, temple_remove: bool = True) -> str:
    """Apply all text replacement fixes. temple_remove=False keeps 'I pressed my temple'."""
    result = text
    for pat, repl in ELLIPSIS_FIXES:
        result = re.sub(pat, repl, result)
    for pat, repl in MARCO_RING_FIXES:
        result = re.sub(pat, repl, result)
    if temple_remove:
        for pat, repl in TEMPLE_REMOVALS:
            result = re.sub(pat, repl, result)
    for pat, repl in RUFFLED_FIXES:
        result = re.sub(pat, repl, result)
    for pat, repl in CONFESSION_VERDICT_FIXES:
        result = re.sub(pat, repl, result, count=1)
    for pat, repl in POV_FIXES:
        result = re.sub(pat, repl, result)
    return result


def apply_paragraph_fixes(doc: Document) -> dict:
    """Apply fixes to each paragraph; remove duplicate block. Returns stats."""
    stats = {"ellipsis": 0, "marco_ring": 0, "temple": 0, "pov": 0, "duplicates_removed": 0}

    # First pass: text replacements
    temple_count = 0
    for para in doc.paragraphs:
        orig = para.text
        # Keep "I pressed my temple" in first 2 paragraphs that have it; remove from rest
        has_temple = "pressed my temple" in orig
        temple_remove = has_temple and temple_count >= 2
        if has_temple:
            temple_count += 1
        new = apply_text_fixes(orig, temple_remove=temple_remove)
        if new != orig:
            para.clear()
            para.add_run(new)
            if "..." in new and "..." not in orig:
                stats["ellipsis"] += 1
            if "different from Marco" in orig and "different from Marco" not in new:
                stats["marco_ring"] += 1
            if "pressed my temple" in orig and "pressed my temple" not in new:
                stats["temple"] += 1
            if any(p in orig for p in ("my voice", "my fingers", "My eyes find mine")) and "her" in new:
                stats["pov"] += 1

    # Second pass: remove duplicate block (exact duplicate paragraphs)
    paras = list(doc.paragraphs)
    seen = {}
    to_remove = []
    for i, p in enumerate(paras):
        t = p.text.strip()
        if len(t) < 50:
            continue
        if t in seen:
            to_remove.append(i)
        else:
            seen[t] = i

    for idx in sorted(to_remove, reverse=True):
        paras[idx]._p.getparent().remove(paras[idx]._p)
        stats["duplicates_removed"] += 1

    return stats


def main():
    docx_arg = sys.argv[1] if len(sys.argv) > 1 else None
    if docx_arg:
        docx_path = Path(docx_arg)
        if not docx_path.is_absolute():
            docx_path = (PROJECT_ROOT.parent / docx_arg).resolve()
    else:
        docx_path = DEFAULT_DOCX

    if not docx_path.exists():
        print(f"[ERROR] File not found: {docx_path}")
        return 1

    out_path = docx_path.parent / (docx_path.stem + "_audit_fixed.docx")
    if "_audit_fixed" in docx_path.stem:
        out_path = docx_path

    print(f"Loading: {docx_path.name}")
    doc = Document(str(docx_path))
    stats = apply_paragraph_fixes(doc)
    doc.save(str(out_path))

    print(f"Saved: {out_path.name}")
    print(f"  Ellipsis fixes: {stats['ellipsis']}")
    print(f"  Marco ring fixes: {stats['marco_ring']}")
    print(f"  Temple tic removals: {stats['temple']}")
    print(f"  POV fixes: {stats['pov']}")
    print(f"  Duplicate blocks removed: {stats['duplicates_removed']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
