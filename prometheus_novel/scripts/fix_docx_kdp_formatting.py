"""Apply KDP-compliant formatting to an existing .docx manuscript.

Preserves all content; fixes page size, margins, styles, and typography
to match Kindle Direct Publishing requirements (6x9 trim, Garamond, etc.).

Usage:
    python -m prometheus_novel.scripts.fix_docx_kdp_formatting [path/to/manuscript.docx]
    python -m prometheus_novel.scripts.fix_docx_kdp_formatting
      (defaults to Burning Vows Final.docx; saves as <name>_KDP.docx)
"""

import sys
from pathlib import Path

if sys.platform == "win32":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = PROJECT_ROOT / "data/projects/burning-vows-30k/output/Burning Vows Final.docx"

# KDP settings (match docx_exporter.py)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
MARGIN_TOP = Inches(0.75)
MARGIN_BOTTOM = Inches(0.75)
MARGIN_LEFT = Inches(0.75)
MARGIN_RIGHT = Inches(0.5)
MARGIN_GUTTER = Inches(0.25)
BODY_FONT = "Garamond"
BODY_SIZE = Pt(11)
CHAPTER_FONT = "Garamond"
CHAPTER_SIZE = Pt(24)


def apply_kdp_formatting(doc: Document) -> None:
    """Apply KDP page setup and styles to document."""
    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.gutter = MARGIN_GUTTER

    # Normal / body style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Inches(0.3)

    # Heading 1 – chapter titles (and possibly book title)
    # Note: page_break_before only on "Chapter N" to avoid blank first page
    h1 = styles["Heading 1"]
    h1.font.name = CHAPTER_FONT
    h1.font.size = CHAPTER_SIZE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(72)
    h1.paragraph_format.space_after = Pt(24)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.page_break_before = False  # Apply per-chapter below

    # Heading 3 (e.g. "A Romance Novel" on title page)
    if "Heading 3" in styles:
        h3 = styles["Heading 3"]
        h3.font.name = CHAPTER_FONT
        h3.font.size = Pt(14)
        h3.font.italic = True
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h3.paragraph_format.first_line_indent = Inches(0)

    # Scene Break style for * * *
    try:
        scene_break = styles.add_style("Scene Break", WD_STYLE_TYPE.PARAGRAPH)
        scene_break.font.name = BODY_FONT
        scene_break.font.size = BODY_SIZE
        scene_break.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scene_break.paragraph_format.space_before = Pt(12)
        scene_break.paragraph_format.space_after = Pt(12)
        scene_break.paragraph_format.first_line_indent = Inches(0)
    except ValueError:
        scene_break = styles["Scene Break"]
        scene_break.font.name = BODY_FONT
        scene_break.font.size = BODY_SIZE
        scene_break.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scene_break.paragraph_format.space_before = Pt(12)
        scene_break.paragraph_format.space_after = Pt(12)
        scene_break.paragraph_format.first_line_indent = Inches(0)

    # First paragraph after each chapter: no indent
    try:
        first_para = styles.add_style("First Paragraph", WD_STYLE_TYPE.PARAGRAPH)
        first_para.base_style = normal
        first_para.font.name = BODY_FONT
        first_para.font.size = BODY_SIZE
        first_para.paragraph_format.first_line_indent = Inches(0)
        first_para.paragraph_format.space_after = Pt(0)
        first_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except ValueError:
        first_para = styles["First Paragraph"]
        first_para.font.name = BODY_FONT
        first_para.font.size = BODY_SIZE
        first_para.paragraph_format.first_line_indent = Inches(0)

    # Apply Scene Break to * * * paragraphs
    for para in doc.paragraphs:
        if para.text.strip() == "* * *":
            para.style = "Scene Break"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply page_break_before only to "Chapter N" headings (not title "BURNING VOWS")
    # and first paragraph after each chapter: no indent
    prev_was_chapter = False
    for para in doc.paragraphs:
        text = para.text.strip()
        is_chapter = para.style.name == "Heading 1" and text.startswith("Chapter ")
        if is_chapter:
            para.paragraph_format.page_break_before = True
            prev_was_chapter = True
            continue
        if prev_was_chapter and para.style.name == "Normal" and text:
            para.style = "First Paragraph"
        prev_was_chapter = False


def main():
    docx_arg = sys.argv[1] if len(sys.argv) > 1 else None
    if docx_arg:
        docx_path = Path(docx_arg)
        if not docx_path.is_absolute():
            docx_path = (PROJECT_ROOT.parent / docx_arg).resolve()
    else:
        docx_path = DEFAULT_DOCX

    if not docx_path.exists():
        print(f"[ERROR] File not found: {docx_path}")
        return 1

    # Output: same name with _KDP before .docx
    stem = docx_path.stem
    if stem.endswith("_KDP"):
        out_path = docx_path  # overwrite existing KDP
    else:
        out_path = docx_path.parent / (stem + "_KDP.docx")

    print(f"Loading: {docx_path.name}")
    doc = Document(str(docx_path))
    apply_kdp_formatting(doc)
    doc.save(str(out_path))
    print(f"Saved: {out_path.name}")
    print("KDP formatting applied: 6x9, margins, Garamond 11pt, first-line indent, chapter page breaks")
    return 0


if __name__ == "__main__":
    sys.exit(main())
