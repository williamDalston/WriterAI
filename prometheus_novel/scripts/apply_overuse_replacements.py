"""Apply phrase replacements from overuse config to a .docx manuscript.

Uses phrase_replacement_banks and/or overuse_detected.yaml. Only applies
phrases that have replacements defined (skips empty []). Keeps first N
occurrences, rotates replacements for the rest.

Usage:
    python -m prometheus_novel.scripts.apply_overuse_replacements [docx_path]
    python -m prometheus_novel.scripts.apply_overuse_replacements --docx path.docx --config overuse_detected.yaml
"""

import re
import sys
from pathlib import Path

if sys.platform == "win32":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
    import yaml
except ImportError:
    print("[ERROR] pip install python-docx pyyaml")
    sys.exit(1)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = PROJECT_ROOT / "data/projects/burning-vows-30k/output/Burning Vows Final_KDP_audit_fixed.docx"


def load_banks(*paths: Path) -> dict:
    """Load and merge replacement banks from YAML files."""
    banks = {}
    for p in paths:
        if not p.exists():
            continue
        try:
            with open(p, encoding="utf-8") as f:
                data = yaml.safe_load(f) or {}
            raw = data.get("banks") or data
            if isinstance(raw, dict):
                for k, v in raw.items():
                    if isinstance(v, list) and v:
                        banks[str(k).strip()] = [str(x) for x in v]
        except Exception as e:
            print(f"[WARN] Skip {p}: {e}")
    return banks


def apply_replacements(doc: Document, banks: dict, keep_first: int = 2) -> dict:
    """Apply phrase replacements to doc paragraphs. Returns stats."""
    stats = {}
    for phrase, replacements in banks.items():
        if not replacements:
            continue
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        ridx = 0
        count = 0
        replaced = 0
        for para in doc.paragraphs:
            text = para.text
            matches = list(pattern.finditer(text))
            if not matches:
                continue
            new_text = text
            for m in reversed(matches):
                count += 1
                if count <= keep_first:
                    continue
                repl = replacements[ridx % len(replacements)]
                ridx += 1
                orig = m.group()
                if repl and orig and orig[0].isupper():
                    repl = repl[0].upper() + repl[1:]
                new_text = new_text[: m.start()] + repl + new_text[m.end() :]
                replaced += 1
            if new_text != text:
                para.clear()
                para.add_run(new_text)
        if count > 0:
            stats[phrase] = {"found": count, "replaced": replaced}
    return stats


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("path", nargs="?", help="Docx path")
    parser.add_argument("--docx", "-d", help="Docx path")
    parser.add_argument("--config", "-c", nargs="+", help="YAML config(s): banks or overuse_detected")
    parser.add_argument("--keep-first", "-k", type=int, default=2, help="Keep first N (default: 2)")
    args = parser.parse_args()

    docx_path = args.docx or args.path
    if not docx_path:
        docx_path = str(DEFAULT_DOCX)
    docx_path = Path(docx_path)
    if not docx_path.is_absolute():
        docx_path = (PROJECT_ROOT.parent / docx_path).resolve()

    if not docx_path.exists():
        print(f"[ERROR] Not found: {docx_path}")
        return 1

    config_paths = args.config or []
    proj_dir = docx_path.parent
    default_configs = [
        PROJECT_ROOT / "configs" / "phrase_replacement_banks.yaml",
        proj_dir / "phrase_replacement_banks.yaml",
        proj_dir / "overuse_detected.yaml",
    ]
    paths = [Path(p).resolve() for p in config_paths] if config_paths else default_configs

    banks = load_banks(*paths)
    if not banks:
        print("[WARN] No phrases with replacements found. Add alternatives to overuse_detected.yaml or phrase_replacement_banks.yaml")
        return 0

    print(f"Loading: {docx_path.name}")
    doc = Document(str(docx_path))
    stats = apply_replacements(doc, banks, keep_first=args.keep_first)

    out_path = docx_path.parent / (docx_path.stem + "_replaced.docx")
    if "_replaced" in docx_path.stem:
        out_path = docx_path
    doc.save(str(out_path))

    print(f"Saved: {out_path.name}")
    total = sum(s["replaced"] for s in stats.values())
    print(f"Replacements: {total} across {len(stats)} phrases")
    for phrase, s in list(stats.items())[:10]:
        if s["replaced"] > 0:
            print(f"  \"{phrase}\": {s['replaced']}/{s['found']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
