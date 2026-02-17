"""
Word Document Exporter for Kindle Direct Publishing (KDP)

Creates professionally formatted .docx files ready for upload to
Amazon Kindle Direct Publishing.
"""

import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
import json
import yaml

from .scene_validator import validate_project_scenes, _validation_mode

logger = logging.getLogger(__name__)

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT


class KDPExporter:
    """Export novel to Word document formatted for Kindle Direct Publishing."""

    # KDP recommended settings for 6x9 trim size
    PAGE_WIDTH = Inches(6)
    PAGE_HEIGHT = Inches(9)
    MARGIN_TOP = Inches(0.75)
    MARGIN_BOTTOM = Inches(0.75)
    MARGIN_LEFT = Inches(0.75)
    MARGIN_RIGHT = Inches(0.5)
    MARGIN_GUTTER = Inches(0.25)

    # Typography
    BODY_FONT = "Garamond"
    BODY_SIZE = Pt(11)
    CHAPTER_FONT = "Garamond"
    CHAPTER_SIZE = Pt(24)

    def __init__(self, project_path: Path):
        self.project_path = Path(project_path)
        self.config: Dict[str, Any] = {}
        self.scenes: List[Dict[str, Any]] = []
        self.doc: Optional[Document] = None

    def load_project(self):
        """Load project config and generated content."""
        # Load config
        config_file = self.project_path / "config.yaml"
        if config_file.exists():
            # Try UTF-8 first, then fallback to latin-1 for special characters
            try:
                with open(config_file, 'r', encoding='utf-8') as f:
                    self.config = yaml.safe_load(f) or {}
            except UnicodeDecodeError:
                with open(config_file, 'r', encoding='latin-1') as f:
                    self.config = yaml.safe_load(f) or {}

        # Load pipeline state if exists
        state_file = self.project_path / "pipeline_state.json"
        if state_file.exists():
            try:
                with open(state_file, 'r', encoding='utf-8') as f:
                    state = json.load(f)
                    self.scenes = state.get("scenes", [])
            except (json.JSONDecodeError, OSError) as e:
                logger.warning(f"Failed to load pipeline state: {e}")
                self.scenes = []

        # Or load from markdown output
        output_file = self.project_path / "output" / f"{self.config.get('project_name', 'novel')}.md"
        if output_file.exists() and not self.scenes:
            self._parse_markdown_output(output_file)

    def _parse_markdown_output(self, md_file: Path):
        """Parse markdown output file into scenes."""
        content = md_file.read_text(encoding='utf-8')

        # Split by chapter headers
        chapters = re.split(r'\n## Chapter (\d+)', content)

        current_chapter = 0
        for i, chunk in enumerate(chapters):
            if chunk.strip().isdigit():
                current_chapter = int(chunk)
            elif chunk.strip():
                # This is content
                scene_parts = re.split(r', Scene (\d+)\n', chunk)
                scene_num = 0
                for j, part in enumerate(scene_parts):
                    if part.strip().isdigit():
                        scene_num = int(part)
                    elif part.strip():
                        self.scenes.append({
                            "chapter": current_chapter,
                            "scene_number": scene_num or 1,
                            "content": part.strip()
                        })

    def _setup_document(self):
        """Create and configure document with KDP settings."""
        self.doc = Document()

        # Set up page size and margins for 6x9 book
        section = self.doc.sections[0]
        section.page_width = self.PAGE_WIDTH
        section.page_height = self.PAGE_HEIGHT
        section.top_margin = self.MARGIN_TOP
        section.bottom_margin = self.MARGIN_BOTTOM
        section.left_margin = self.MARGIN_LEFT
        section.right_margin = self.MARGIN_RIGHT
        section.gutter = self.MARGIN_GUTTER

        # Set up styles
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for professional typography."""
        styles = self.doc.styles

        # Normal/Body style
        normal = styles['Normal']
        normal.font.name = self.BODY_FONT
        normal.font.size = self.BODY_SIZE
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.first_line_indent = Inches(0.3)

        # Heading 1 - Chapter titles
        heading1 = styles['Heading 1']
        heading1.font.name = self.CHAPTER_FONT
        heading1.font.size = self.CHAPTER_SIZE
        heading1.font.bold = True
        heading1.paragraph_format.space_before = Pt(72)  # Start lower on page
        heading1.paragraph_format.space_after = Pt(24)
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading1.paragraph_format.first_line_indent = Inches(0)
        heading1.paragraph_format.page_break_before = True

        # Create style for first paragraph (no indent)
        try:
            first_para = styles.add_style('First Paragraph', WD_STYLE_TYPE.PARAGRAPH)
            first_para.base_style = normal
            first_para.font.name = self.BODY_FONT
            first_para.font.size = self.BODY_SIZE
            first_para.paragraph_format.first_line_indent = Inches(0)
        except ValueError:
            pass  # Style already exists

        # Scene break style
        try:
            scene_break = styles.add_style('Scene Break', WD_STYLE_TYPE.PARAGRAPH)
            scene_break.font.name = self.BODY_FONT
            scene_break.font.size = self.BODY_SIZE
            scene_break.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            scene_break.paragraph_format.space_before = Pt(12)
            scene_break.paragraph_format.space_after = Pt(12)
            scene_break.paragraph_format.first_line_indent = Inches(0)
        except ValueError:
            pass

    def _add_title_page(self):
        """Add title page."""
        title = self.config.get("title", "Untitled Novel")

        # Add blank space at top
        for _ in range(8):
            self.doc.add_paragraph()

        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title.upper())
        run.font.name = self.CHAPTER_FONT
        run.font.size = Pt(28)
        run.font.bold = True

        # Subtitle/genre if exists
        if self.config.get("genre"):
            self.doc.add_paragraph()
            subtitle = self.doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run(f"A {self.config.get('genre', '').title()} Novel")
            run.font.name = self.CHAPTER_FONT
            run.font.size = Pt(14)
            run.font.italic = True

        # Add page break after title page
        self.doc.add_page_break()

    def _add_copyright_page(self):
        """Add copyright page."""
        title = self.config.get("title", "Untitled Novel")

        # Blank space
        for _ in range(20):
            self.doc.add_paragraph()

        # Copyright notice
        copyright_para = self.doc.add_paragraph()
        copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_para.paragraph_format.first_line_indent = Inches(0)

        lines = [
            f"{title}",
            "",
            "Copyright (c) 2026",
            "All rights reserved.",
            "",
            "This is a work of fiction. Names, characters, places, and incidents",
            "either are products of the author's imagination or are used fictitiously.",
            "Any resemblance to actual events or locales or persons,",
            "living or dead, is entirely coincidental.",
            "",
            "No part of this book may be reproduced in any form or by any electronic",
            "or mechanical means, including information storage and retrieval systems,",
            "without written permission from the author, except for the use of",
            "brief quotations in a book review."
        ]

        for line in lines:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Inches(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(line)
            run.font.name = self.BODY_FONT
            run.font.size = Pt(10)

        self.doc.add_page_break()

    def _add_formatted_paragraph(self, text: str, first_in_section: bool = False):
        """Add a paragraph with Markdown bold/italic rendered as DOCX formatting.

        Parses ``**bold**`` and ``*italic*`` markers into proper DOCX runs.
        Bold markers are matched first so ``**`` is not misinterpreted as
        two italic markers.  Unmatched asterisks are left as-is.
        """
        para = self.doc.add_paragraph()
        if first_in_section:
            para.paragraph_format.first_line_indent = Inches(0)

        # Pattern: **bold** (group 2), then *italic* (group 3).
        # Order in the alternation matters: **...** is tried before *...*
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Add any plain text before this match
            if match.start() > last_end:
                para.add_run(text[last_end:match.start()])

            if match.group(2):  # **bold**
                run = para.add_run(match.group(2))
                run.font.bold = True
            elif match.group(3):  # *italic*
                run = para.add_run(match.group(3))
                run.font.italic = True

            last_end = match.end()

        # Add any remaining plain text after the last match
        if last_end < len(text):
            para.add_run(text[last_end:])

        return para

    def _add_chapter(self, chapter_num: int, scenes: List[Dict[str, Any]]):
        """Add a chapter with its scenes."""
        # Chapter heading
        chapter_title = f"Chapter {chapter_num}"
        heading = self.doc.add_heading(chapter_title, level=1)

        # Process scenes
        for i, scene in enumerate(scenes):
            content = scene.get("content", "")

            # Add scene break between scenes (except first)
            if i > 0:
                break_para = self.doc.add_paragraph("* * *")
                break_para.style = 'Scene Break' if 'Scene Break' in self.doc.styles else 'Normal'
                break_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Split content into paragraphs
            paragraphs = content.split('\n\n')

            for j, para_text in enumerate(paragraphs):
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Clean up the text
                para_text = para_text.replace('\n', ' ')

                para = self._add_formatted_paragraph(para_text, first_in_section=(j == 0))

    def _validate_scenes(self) -> None:
        """Run pre-export validation; log issues. In strict mode, raise on errors."""
        if not self.scenes:
            return
        report = validate_project_scenes(self.scenes, self.config)
        for i in report["issues"]:
            msg = f"[{i['code']}] {i['scene_id']}: {i['message']}"
            if i.get("excerpt"):
                msg += f" | excerpt: {i['excerpt'][:80]}"
            if i["severity"] == "error":
                logger.error(msg)
            else:
                logger.warning(msg)
        mode = _validation_mode(self.config)
        if report["has_errors"] and mode == "strict":
            raise ValueError(
                f"Export blocked: {len([x for x in report['issues'] if x['severity'] == 'error'])} "
                "validation error(s). Set export.validation_mode: lenient to allow export."
            )

    def export(self, output_path: Optional[Path] = None) -> Path:
        """Export the novel to a Word document.

        Args:
            output_path: Custom output path. If None, saves to project output folder.

        Returns:
            Path to the generated .docx file.
        """
        self.load_project()
        self._validate_scenes()
        self._setup_document()

        # Add front matter
        self._add_title_page()
        self._add_copyright_page()

        # Group scenes by chapter (skip malformed entries)
        chapters: Dict[int, List[Dict]] = {}
        for scene in self.scenes:
            if not isinstance(scene, dict):
                continue
            ch = scene.get("chapter", 1)
            if ch not in chapters:
                chapters[ch] = []
            chapters[ch].append(scene)

        # Add chapters in order, scenes sorted within each chapter
        def _scene_sort_key(sc: dict):
            return (
                int(sc.get("scene_number") or sc.get("scene") or 10**9),
                str(sc.get("scene_id") or ""),
            )

        for chapter_num in sorted(chapters.keys()):
            self._add_chapter(chapter_num, sorted(chapters[chapter_num], key=_scene_sort_key))

        # Determine output path
        if output_path is None:
            output_dir = self.project_path / "output"
            output_dir.mkdir(exist_ok=True)
            project_name = self.config.get("project_name", "novel")
            output_path = output_dir / f"{project_name}_KDP.docx"

        # Save document
        self.doc.save(str(output_path))

        return output_path

    def export_sample(self, output_path: Optional[Path] = None) -> Path:
        """Export a sample document from config data (no generation needed).

        Creates a Word doc with title page, copyright, and seed content
        formatted as a preview/template.
        """
        self.load_project()
        self._setup_document()

        # Add front matter
        self._add_title_page()
        self._add_copyright_page()

        # Add synopsis as Chapter 1 placeholder
        synopsis = self.config.get("synopsis", "")
        if synopsis:
            heading = self.doc.add_heading("Chapter 1", level=1)

            # Add synopsis content
            for para_text in synopsis.split('\n\n'):
                if para_text.strip():
                    para = self.doc.add_paragraph(para_text.strip())

        # Add placeholder chapters based on key plot points
        key_points = self.config.get("key_plot_points", "")
        if key_points:
            heading = self.doc.add_heading("Story Outline", level=1)
            for para_text in key_points.split('\n'):
                if para_text.strip():
                    para = self.doc.add_paragraph(para_text.strip())

        # Add character info
        if self.config.get("protagonist") or self.config.get("antagonist"):
            heading = self.doc.add_heading("Character Notes", level=1)

            if self.config.get("protagonist"):
                para = self.doc.add_paragraph()
                para.add_run("Protagonist: ").bold = True
                para.add_run(self.config.get("protagonist"))

            if self.config.get("antagonist"):
                para = self.doc.add_paragraph()
                para.add_run("Antagonist: ").bold = True
                para.add_run(self.config.get("antagonist"))

        # Determine output path
        if output_path is None:
            output_dir = self.project_path / "output"
            output_dir.mkdir(exist_ok=True)
            project_name = self.config.get("project_name", "novel")
            output_path = output_dir / f"{project_name}_seed.docx"

        self.doc.save(str(output_path))
        return output_path


def export_to_docx(project_path: str, sample_only: bool = False) -> str:
    """Convenience function to export project to Word document.

    Args:
        project_path: Path to project directory
        sample_only: If True, export seed data only (no generation needed)

    Returns:
        Path to generated .docx file
    """
    exporter = KDPExporter(Path(project_path))
    if sample_only:
        return str(exporter.export_sample())
    return str(exporter.export())
