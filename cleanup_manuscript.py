#!/usr/bin/env python3
"""Post-hoc cleanup of a generated manuscript .docx file.

Removes all known AI artifact categories:
1. LLM preambles ("Sure, here's a revised version...")
2. Editing notes ("Changes made:", "Scanning for AI tells:")
3. Prompt bleed-through ("CURRENT SCENE:", "A great chapter-ending hook can be:")
4. Section markers ("=== EXPANDED SCENE ===", "ENHANCED SCENE:")
5. UI artifacts ("Visible: 0% – 100%")
6. Meta-text ("The rest of the scene remains unchanged")
7. Duplicate paragraphs (near-identical content appearing twice)

Usage:
    python cleanup_manuscript.py <input.docx> [output.docx]

If output is not specified, saves as <input>_cleaned.docx
"""

import re
import sys
from pathlib import Path
from copy import deepcopy

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── Artifact patterns to REMOVE entirely ─────────────────────────────────────

# Patterns that match entire paragraphs to delete
PARAGRAPH_DELETE_PATTERNS = [
    # LLM preambles
    r'^(?:Sure[,!.]?\s*)?[Hh]ere\'?s?\s+(?:the |a |my |an? )?'
    r'(?:revised|enhanced|polished|expanded|edited|updated|improved|rewritten|final)\b',
    r'^(?:Sure|Certainly|Of course|Absolutely)[,!.]\s*(?:here|I\'ve|I have|let me)\b',
    r'^(?:I\'ve |I have )(?:revised|enhanced|polished|expanded|edited|rewritten)\b',
    r'^(?:Below is|The following is|What follows is)\b',
    # Meta-text / editing notes
    r'(?:The )?rest (?:of (?:the |this )?(?:scene|chapter|text|content) )?(?:remains?|is) unchanged',
    r'^\[The rest (?:of (?:the |this )?(?:scene|chapter))? remains unchanged',
    r'^(?:CURRENT SCENE(?: MODIFIED)?:)',
    r'^(?:ENHANCED SCENE|EXPANDED SCENE|POLISHED SCENE|FIXED SCENE|REVISED SCENE):',
    r'^(?:Changes made|Scanning for AI|Quality checklist|AI tells)',
    r'^(?:\*\*(?:Changes|Scanning|Notes?|Quality|Summary|Checklist))',
    r'^(?:Output (?:ONLY |only )?the (?:revised|enhanced|polished))',
    # Prompt bleed-through
    r'^A great chapter-(?:ending|opening) hook can be:',
    r'^[-•*]\s*(?:A cliffhanger|In medias res|A striking sensory|A provocative)',
    r'^[-•*]\s*(?:Immediate conflict|Disorientation|A kiss or romantic)',
    r'^[-•*]\s*(?:A threat delivered|A question raised|A twist revealed)',
    r'^[-•*]\s*(?:An emotional gut-punch|A decision made)',
    # UI artifacts
    r'^Visible:\s*\d+%',
    # Section markers
    r'^---+$',
    r'^\*\*\*+$',
    r'^===+.*===*$',
    r'^#{1,3}\s+(?:Version|Take|Alternative|Revised|Enhanced)',
]

# Patterns to remove INLINE (within a paragraph, replace with empty)
INLINE_STRIP_PATTERNS = [
    r'(?:The )?rest (?:of (?:the |this )?(?:scene|chapter|text|content) )?(?:remains?|is) unchanged[^.]*[.\s]*',
    r'\[The rest (?:of (?:the |this )?(?:scene|chapter))? remains unchanged[^\]]*\]\s*',
    r'CURRENT SCENE(?: MODIFIED)?:\s*',
    r'Visible:\s*\d+%\s*[–—-]\s*\d+%\s*',
    r'(?:ENHANCED|EXPANDED|POLISHED|FIXED|REVISED) SCENE:\s*',
    r'(?:Sure[,.]?\s*)?[Hh]ere\'?s?\s+(?:the |a )?'
    r'(?:revised|enhanced|polished|expanded|edited)\s+'
    r'(?:version|scene|text|content)[.:]\s*',
]

# Compile all patterns
COMPILED_DELETE = [re.compile(p, re.IGNORECASE) for p in PARAGRAPH_DELETE_PATTERNS]
COMPILED_INLINE = [re.compile(p, re.IGNORECASE) for p in INLINE_STRIP_PATTERNS]


def should_delete_paragraph(text: str) -> bool:
    """Check if an entire paragraph should be deleted."""
    stripped = text.strip()
    if not stripped:
        return False
    for pattern in COMPILED_DELETE:
        if pattern.search(stripped):
            return True
    return False


def clean_paragraph_text(text: str) -> str:
    """Clean inline artifacts from paragraph text."""
    for pattern in COMPILED_INLINE:
        text = pattern.sub('', text)
    # Clean up whitespace
    text = re.sub(r'  +', ' ', text)
    return text.strip()


def find_duplicate_paragraphs(paragraphs):
    """Find indices of near-duplicate paragraphs to remove."""
    duplicates = set()
    fingerprints = {}

    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if len(text) < 80:
            continue

        words = text.split()
        if len(words) < 10:
            continue

        fingerprint = ' '.join(words[:8]).lower()

        if fingerprint in fingerprints:
            orig_idx = fingerprints[fingerprint]
            orig_text = paragraphs[orig_idx].text.strip()
            orig_words = set(orig_text.lower().split())
            curr_words = set(text.lower().split())
            overlap = len(orig_words & curr_words) / min(len(orig_words), len(curr_words))
            if overlap > 0.6:
                duplicates.add(idx)
        else:
            fingerprints[fingerprint] = idx

    return duplicates


def cleanup_docx(input_path: str, output_path: str = None):
    """Clean a manuscript .docx file of all known AI artifacts."""
    doc = Document(input_path)

    if not output_path:
        p = Path(input_path)
        output_path = str(p.parent / f"{p.stem}_cleaned{p.suffix}")

    stats = {
        "paragraphs_deleted": 0,
        "paragraphs_cleaned": 0,
        "duplicates_removed": 0,
        "total_paragraphs": len(doc.paragraphs),
    }

    # Phase 1: Find duplicate paragraphs
    duplicate_indices = find_duplicate_paragraphs(doc.paragraphs)
    stats["duplicates_removed"] = len(duplicate_indices)

    # Phase 2: Process paragraphs (mark for deletion or clean inline)
    paragraphs_to_delete = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text

        # Check if this is a duplicate
        if idx in duplicate_indices:
            paragraphs_to_delete.append(para)
            stats["paragraphs_deleted"] += 1
            continue

        # Check if entire paragraph should be deleted
        if should_delete_paragraph(text):
            paragraphs_to_delete.append(para)
            stats["paragraphs_deleted"] += 1
            continue

        # Clean inline artifacts
        cleaned = clean_paragraph_text(text)
        if cleaned != text:
            stats["paragraphs_cleaned"] += 1
            # Update paragraph text while preserving formatting
            if para.runs:
                # Put cleaned text in first run, clear the rest
                para.runs[0].text = cleaned
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = cleaned

    # Phase 3: Delete marked paragraphs
    for para in paragraphs_to_delete:
        p = para._element
        p.getparent().remove(p)

    # Save
    doc.save(output_path)

    return output_path, stats


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not Path(input_path).exists():
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    print(f"Cleaning: {input_path}")
    output, stats = cleanup_docx(input_path, output_path)

    print(f"\nResults:")
    print(f"  Total paragraphs scanned: {stats['total_paragraphs']}")
    print(f"  Paragraphs deleted:       {stats['paragraphs_deleted']}")
    print(f"  Paragraphs cleaned:       {stats['paragraphs_cleaned']}")
    print(f"  Duplicates removed:       {stats['duplicates_removed']}")
    print(f"\nSaved to: {output}")


if __name__ == "__main__":
    main()
